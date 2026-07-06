from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

doc = Document()

# ── Estilos por defecto ──
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hs.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)

# ── Helper para código ──
def add_code_block(doc, code, language=''):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    run = p.add_run(code)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    return p

def add_table_row(table, cells_data, bold=False):
    row = table.add_row()
    for i, text in enumerate(cells_data):
        cell = row.cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(str(text))
        run.font.size = Pt(9)
        if bold:
            run.bold = True
        # Shade header
        if bold:
            shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="2C3E50"/>')
            cell._tc.get_or_add_tcPr().append(shading)
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    return row

def add_simple_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="2C3E50"/>')
        cell._tc.get_or_add_tcPr().append(shading)
    for row_data in rows:
        add_table_row(table, row_data)
    doc.add_paragraph()
    return table

# ══════════════════════════════════════════════════════════════
# PORTADA
# ══════════════════════════════════════════════════════════════
for _ in range(6):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('MANUAL TÉCNICO')
run.bold = True
run.font.size = Pt(28)
run.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Compra tu Coquito')
run.bold = True
run.font.size = Pt(36)
run.font.color.rgb = RGBColor(0xC1, 0x7F, 0x3E)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Sistema de Punto de Venta (POS)')
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(0x7F, 0x8C, 0x8D)

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Versión 2.0 — Julio 2026')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x95, 0xA5, 0xA6)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Stack: XAMPP | PHP 8 | MySQL | HTML5 | CSS3 | JavaScript')
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x95, 0xA5, 0xA6)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# ÍNDICE
# ══════════════════════════════════════════════════════════════
doc.add_heading('Índice', level=1)
toc_items = [
    '1. Introducción',
    '2. Arquitectura del Sistema',
    '3. Estructura de Directorios',
    '4. Base de Datos',
    '   4.1 Modelo Entidad-Relación',
    '   4.2 Diccionario de Tablas',
    '   4.3 Relaciones',
    '5. Backend (PHP)',
    '   5.1 Configuración',
    '   5.2 Modelos',
    '   5.3 Controladores y API REST',
    '   5.4 Autenticación',
    '6. Frontend',
    '   6.1 HTML (index.html)',
    '   6.2 CSS (style.css)',
    '   6.3 JavaScript (script.js)',
    '   6.4 Flujo de la Aplicación',
    '7. API Endpoints — Referencia Rápida',
    '8. Instalación y Despliegue',
    '9. Guía de Uso',
    '10. Mantenimiento y Mejoras',
    'Anexo A — Script SQL Completo',
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(0)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 1. INTRODUCCIÓN
# ══════════════════════════════════════════════════════════════
doc.add_heading('1. Introducción', level=1)
doc.add_paragraph(
    'Compra tu Coquito es un sistema de punto de venta (POS) web desarrollado para '
    'una librería y papelería. Permite gestionar el catálogo de productos, registrar '
    'ventas con cálculo automático de IGV (18%), y consultar el historial de transacciones.'
)
doc.add_paragraph('El sistema sigue una arquitectura cliente-servidor tradicional:')
bullets = [
    'Frontend: Single Page Application (SPA) con HTML, CSS y JavaScript vanilla.',
    'Backend: API REST en PHP 8 sin framework.',
    'Base de datos: MySQL relacional con 7 tablas.',
    'Servidor: Apache HTTP (XAMPP).',
]
for b in bullets:
    doc.add_paragraph(b, style='List Bullet')

doc.add_paragraph(
    'El proyecto es de código abierto con fines académicos y demostrativos. '
    'Las contraseñas se almacenan con bcrypt y las sesiones se manejan con cookies PHP.'
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 2. ARQUITECTURA DEL SISTEMA
# ══════════════════════════════════════════════════════════════
doc.add_heading('2. Arquitectura del Sistema', level=1)

doc.add_heading('2.1 Stack Tecnológico', level=2)
add_simple_table(doc,
    ['Componente', 'Tecnología', 'Versión'],
    [
        ['Servidor web', 'Apache (XAMPP)', '2.4.x'],
        ['Backend', 'PHP', '8.x'],
        ['Base de datos', 'MySQL (MariaDB)', '10.x'],
        ['Frontend', 'HTML5 / CSS3 / JavaScript', 'ES6'],
        ['Fuentes externas', 'Google Fonts (Playfair Display, DM Sans)', '—'],
        ['Autenticación', 'bcrypt + sesiones PHP', '—'],
    ]
)

doc.add_heading('2.2 Diagrama de Arquitectura', level=2)
doc.add_paragraph(
    'Navegador (SPA) → Apache (index.html + assets) → API PHP (backend/controllers/) '
    '→ MySQL (bd_coquito). El frontend consume los endpoints mediante fetch() AJAX. '
    'Las rutas protegidas validan la sesión antes de ejecutar cualquier acción.'
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 3. ESTRUCTURA DE DIRECTORIOS
# ══════════════════════════════════════════════════════════════
doc.add_heading('3. Estructura de Directorios', level=1)
add_code_block(doc, '''coquito_proyecto/
├── README.md
├── backend/
│   ├── config/
│   │   ├── conexion.php          # Conexión MySQL
│   │   └── sesion.php            # Guard de sesión
│   ├── controllers/
│   │   ├── LoginController.php   # Login/logout/estado
│   │   ├── ProductoController.php# CRUD productos
│   │   └── VentaController.php   # Registro/listado ventas
│   └── models/
│       ├── Usuario.php           # Modelo usuario
│       ├── Producto.php          # Modelo producto
│       └── Venta.php             # Modelo venta (transaccional)
├── database/
│   └── bd_coquito.sql            # Esquema + datos iniciales
├── docs/
│   ├── diagrama-er.png           # Diagrama entidad-relación
│   └── Manual_Tecnico_Compra_tu_Coquito.docx
└── frontend/
    ├── index.html                # SPA (HTML puro)
    ├── css/
    │   └── style.css             # Estilos completos
    ├── js/
    │   └── script.js             # Lógica del cliente
    └── img/                      # Imágenes (3 archivos)''')

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 4. BASE DE DATOS
# ══════════════════════════════════════════════════════════════
doc.add_heading('4. Base de Datos', level=1)

doc.add_heading('4.1 Modelo Entidad-Relación', level=2)
diagram_path = os.path.join(os.path.dirname(__file__), 'diagrama-er.png')
if os.path.exists(diagram_path):
    doc.add_picture(diagram_path, width=Inches(6))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_heading('4.2 Diccionario de Tablas', level=2)

# roles
doc.add_heading('roles', level=3)
doc.add_paragraph('Almacena los roles del sistema (Administrador, Vendedor).')
add_simple_table(doc,
    ['Campo', 'Tipo', 'Restricciones', 'Descripción'],
    [
        ['id_rol', 'INT', 'AUTO_INCREMENT, PK', 'Identificador único'],
        ['nombre_rol', 'VARCHAR(50)', 'NOT NULL', 'Nombre del rol'],
    ]
)

# usuarios
doc.add_heading('usuarios', level=3)
doc.add_paragraph('Usuarios del sistema con autenticación por password bcrypt.')
add_simple_table(doc,
    ['Campo', 'Tipo', 'Restricciones', 'Descripción'],
    [
        ['id_usuario', 'INT', 'AUTO_INCREMENT, PK', 'Identificador único'],
        ['nombres', 'VARCHAR(100)', 'NOT NULL', 'Nombre(s) del usuario'],
        ['apellidos', 'VARCHAR(100)', 'NOT NULL', 'Apellidos del usuario'],
        ['usuario', 'VARCHAR(50)', 'NOT NULL, UNIQUE', 'Nombre de usuario (login)'],
        ['correo', 'VARCHAR(100)', 'NOT NULL, UNIQUE', 'Correo electrónico'],
        ['password', 'VARCHAR(255)', 'NOT NULL', 'Hash bcrypt de la contraseña'],
        ['id_rol', 'INT', 'NOT NULL, FK → roles', 'Rol asignado'],
        ['estado', 'TINYINT', 'DEFAULT 1', '1=activo, 0=inactivo'],
        ['fecha_registro', 'DATETIME', 'DEFAULT CURRENT_TIMESTAMP', 'Fecha de creación'],
    ]
)

# categorias
doc.add_heading('categorias', level=3)
doc.add_paragraph('Catálogo de categorías de productos.')
add_simple_table(doc,
    ['Campo', 'Tipo', 'Restricciones', 'Descripción'],
    [
        ['id_categoria', 'INT', 'AUTO_INCREMENT, PK', 'Identificador único'],
        ['nombre_categoria', 'VARCHAR(100)', 'NOT NULL', 'Nombre de la categoría'],
        ['estado', 'TINYINT', 'DEFAULT 1', '1=activo, 0=inactivo'],
    ]
)

# productos
doc.add_heading('productos', level=3)
doc.add_paragraph('Productos disponibles para la venta.')
add_simple_table(doc,
    ['Campo', 'Tipo', 'Restricciones', 'Descripción'],
    [
        ['id_producto', 'INT', 'AUTO_INCREMENT, PK', 'Identificador único'],
        ['nombre_producto', 'VARCHAR(150)', 'NOT NULL', 'Nombre del producto'],
        ['id_categoria', 'INT', 'NOT NULL, FK → categorias', 'Categoría del producto'],
        ['precio', 'DECIMAL(10,2)', 'NOT NULL', 'Precio unitario en Soles'],
        ['stock', 'INT', 'NOT NULL, DEFAULT 0', 'Cantidad en inventario'],
        ['estado', 'TINYINT', 'DEFAULT 1', '1=activo, 0=eliminado (soft-delete)'],
        ['fecha_registro', 'DATETIME', 'DEFAULT CURRENT_TIMESTAMP', 'Fecha de creación'],
    ]
)

# clientes
doc.add_heading('clientes', level=3)
doc.add_paragraph('Clientes creados automáticamente al registrar una venta.')
add_simple_table(doc,
    ['Campo', 'Tipo', 'Restricciones', 'Descripción'],
    [
        ['id_cliente', 'INT', 'AUTO_INCREMENT, PK', 'Identificador único'],
        ['nombre_cliente', 'VARCHAR(150)', 'NOT NULL', 'Nombre del cliente'],
        ['fecha_registro', 'DATETIME', 'DEFAULT CURRENT_TIMESTAMP', 'Fecha de creación'],
    ]
)

# ventas
doc.add_heading('ventas', level=3)
doc.add_paragraph('Cabecera de cada venta.')
add_simple_table(doc,
    ['Campo', 'Tipo', 'Restricciones', 'Descripción'],
    [
        ['id_venta', 'INT', 'AUTO_INCREMENT, PK', 'Número de venta'],
        ['id_cliente', 'INT', 'NULL, FK → clientes', 'Cliente (opcional)'],
        ['id_usuario', 'INT', 'NOT NULL, FK → usuarios', 'Usuario que registró la venta'],
        ['subtotal', 'DECIMAL(10,2)', 'NOT NULL', 'Suma de subtotales de items'],
        ['igv', 'DECIMAL(10,2)', 'NOT NULL', 'IGV = subtotal × 0.18'],
        ['total', 'DECIMAL(10,2)', 'NOT NULL', 'Total = subtotal + igv'],
        ['fecha_venta', 'DATETIME', 'DEFAULT CURRENT_TIMESTAMP', 'Fecha y hora de la venta'],
    ]
)

# detalle_ventas
doc.add_heading('detalle_ventas', level=3)
doc.add_paragraph('Líneas de detalle de cada venta.')
add_simple_table(doc,
    ['Campo', 'Tipo', 'Restricciones', 'Descripción'],
    [
        ['id_detalle', 'INT', 'AUTO_INCREMENT, PK', 'Identificador único'],
        ['id_venta', 'INT', 'NOT NULL, FK → ventas', 'Venta a la que pertenece'],
        ['id_producto', 'INT', 'NOT NULL, FK → productos', 'Producto vendido'],
        ['cantidad', 'INT', 'NOT NULL', 'Cantidad vendida'],
        ['precio_unitario', 'DECIMAL(10,2)', 'NOT NULL', 'Precio al momento de la venta'],
        ['subtotal_item', 'DECIMAL(10,2)', 'NOT NULL', 'cantidad × precio_unitario'],
    ]
)

doc.add_heading('4.3 Relaciones', level=2)
add_simple_table(doc,
    ['Relación', 'Tipo', 'Descripción'],
    [
        ['roles → usuarios', '1:N', 'Un rol tiene muchos usuarios'],
        ['categorias → productos', '1:N', 'Una categoría tiene muchos productos'],
        ['clientes → ventas', '1:N', 'Un cliente realiza muchas ventas (FK nullable)'],
        ['usuarios → ventas', '1:N', 'Un usuario registra muchas ventas'],
        ['ventas → detalle_ventas', '1:N', 'Una venta se compone de muchos detalles'],
        ['productos → detalle_ventas', '1:N', 'Un producto aparece en muchos detalles'],
    ]
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 5. BACKEND
# ══════════════════════════════════════════════════════════════
doc.add_heading('5. Backend (PHP)', level=1)

doc.add_heading('5.1 Configuración', level=2)

doc.add_heading('conexion.php', level=3)
doc.add_paragraph(
    'Establece la conexión a MySQL mediante mysqli. Las credenciales están '
    'hardcodeadas para el entorno local de XAMPP (root sin contraseña).'
)
add_code_block(doc, '''<?php
$servidor = "localhost";
$usuario  = "root";
$password = "";
$bd       = "bd_coquito";

$conexion = new mysqli($servidor, $usuario, $password, $bd);

if ($conexion->connect_error) {
    die(json_encode(["ok" => false, "mensaje" => "Error de conexión: "
        . $conexion->connect_error]));
}
$conexion->set_charset("utf8");
?>''')

doc.add_heading('sesion.php', level=3)
doc.add_paragraph(
    'Función reutilizable que verifica si existe una sesión activa. '
    'Si no hay sesión, devuelve un JSON con error y detiene la ejecución.'
)
add_code_block(doc, '''<?php
session_start();

function requiereSesion() {
    if (!isset($_SESSION["usuario"])) {
        echo json_encode(["ok" => false, "mensaje" => "Sesión no activa"]);
        exit;
    }
}
?>''')

doc.add_heading('5.2 Modelos', level=2)

doc.add_heading('Usuario.php', level=3)
doc.add_paragraph(
    'Método buscarPorUsuario($usuario): consulta el usuario por nombre de '
    'usuario, incluyendo el rol. Solo devuelve registros activos (estado=1).'
)

doc.add_heading('Producto.php', level=3)
doc.add_paragraph('Métodos principales:')
bullets = [
    'listar($busqueda): lista productos activos con su categoría, filtro LIKE por nombre.',
    'crear($nombre, $idCategoria, $precio, $stock): inserta un nuevo producto.',
    'editar($id, $nombre, $idCategoria, $precio, $stock): actualiza un producto existente.',
    'eliminar($id): soft-delete (SET estado=0), no borra físicamente.',
    'descontarStock($id, $cantidad): reduce el stock validando que haya suficiente.',
    'obtenerPorId($id): devuelve un producto por su ID.',
]
for b in bullets:
    doc.add_paragraph(b, style='List Bullet')

doc.add_heading('Venta.php', level=3)
doc.add_paragraph(
    'Modelo transaccional que opera dentro de una transacción MySQL '
    '(begin_transaction / commit / rollback).'
)
bullets = [
    'registrar($idUsuario, $nombreCliente, $items): crea o reusa el cliente, inserta la cabecera de venta, recorre los items descontando stock e insertando detalle, todo en una transacción.',
    'listarRecientes($limite): devuelve las últimas ventas con nombre del cliente y conteo de items.',
    'IGV hardcodeado al 18% (subtotal × 0.18).',
]
for b in bullets:
    doc.add_paragraph(b, style='List Bullet')

doc.add_heading('5.3 Controladores y API REST', level=2)
doc.add_paragraph(
    'Cada controlador recibe una acción vía GET (?accion=...). Todos devuelven JSON. '
    'ProductoController y VentaController llaman a requiereSesion() al inicio.'
)

doc.add_heading('LoginController.php', level=3)
add_simple_table(doc,
    ['Acción', 'Método', 'Descripción'],
    [
        ['login', 'POST', 'Valida usuario/contraseña contra BD, crea sesión'],
        ['logout', 'POST', 'Destruye la sesión'],
        ['estado', 'GET', 'Verifica si hay sesión activa'],
    ]
)

doc.add_heading('ProductoController.php', level=3)
add_simple_table(doc,
    ['Acción', 'Método', 'Parámetros', 'Descripción'],
    [
        ['listar', 'GET', 'busqueda (opcional)', 'Lista productos activos'],
        ['crear', 'POST', 'nombre, id_categoria, precio, stock', 'Crea producto'],
        ['editar', 'POST', 'id, nombre, id_categoria, precio, stock', 'Actualiza producto'],
        ['eliminar', 'POST', 'id', 'Soft-delete de producto'],
    ]
)

doc.add_heading('VentaController.php', level=3)
add_simple_table(doc,
    ['Acción', 'Método', 'Parámetros', 'Descripción'],
    [
        ['registrar', 'POST', 'JSON {cliente, items[]}', 'Registra venta (transaccional)'],
        ['listar', 'GET', '—', 'Lista últimas 20 ventas'],
    ]
)

doc.add_heading('5.4 Autenticación', level=2)
doc.add_paragraph(
    'El sistema usa sesiones PHP nativas (session_start()). El flujo es:'
)
steps = [
    'El usuario envía credenciales a LoginController.php?accion=login.',
    'El backend consulta usuarios por nombre de usuario. Si existe y la contraseña '
    'coincide (password_verify con bcrypt), se crea $_SESSION["usuario"].',
    'Las acciones protegidas (productos, ventas) ejecutan requiereSesion() que '
    'verifica $_SESSION["usuario"] antes de procesar la solicitud.',
    'El frontend verifica la sesión al cargar la página mediante '
    'LoginController.php?accion=estado.',
]
for s in steps:
    doc.add_paragraph(s, style='List Number')

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 6. FRONTEND
# ══════════════════════════════════════════════════════════════
doc.add_heading('6. Frontend', level=1)

doc.add_heading('6.1 HTML (index.html)', level=2)
doc.add_paragraph(
    'Única página HTML que funciona como SPA. Contiene dos vistas principales '
    'que se muestran/ocultan condicionalmente:'
)
bullets = [
    '#pagina-login: portada visual tipo librería con formulario de inicio de sesión, '
    'efectos decorativos (estantería de libros, tarjeta 3D interactiva) y modal de credenciales demo.',
    '#pagina-ventas: panel de ventas con dos columnas (productos + carrito), '
    'más la sección de historial al final.',
    'Carga los assets CSS y JS mediante etiquetas <link> y <script>.',
]
for b in bullets:
    doc.add_paragraph(b, style='List Bullet')

doc.add_heading('6.2 CSS (style.css)', level=2)
doc.add_paragraph(
    '633 líneas de CSS vanilla organizadas en secciones:'
)
add_simple_table(doc,
    ['Sección', 'Líneas', 'Descripción'],
    [
        ['Reset y body', '1-8', 'Estilos base'],
        ['Login / Portada', '9-166', 'Fondo, estantería, tipografía, formulario'],
        ['Panel ventas', '167-206', 'Header y contenedor principal'],
        ['Productos', '207-261', 'Tabla de productos con botones'],
        ['Venta actual', '262-370', 'Carrito, campos, totales, botones'],
        ['Historial', '371-394', 'Tabla de últimas ventas'],
        ['Modal', '395-513', 'Modal de credenciales demo'],
        ['Elementos interactivos', '514-633', 'WhatsApp flotante, tarjeta 3D, grid responsive'],
    ]
)

doc.add_paragraph(
    'Incluye animaciones CSS (wa-pulse para el botón de WhatsApp), '
    'media queries para responsive design, y efecto de glassmorphism '
    'en la portada.'
)

doc.add_heading('6.3 JavaScript (script.js)', level=2)
doc.add_paragraph(
    '353 líneas de JavaScript ES6 vanilla. Se organiza en:'
)
add_simple_table(doc,
    ['Sección', 'Líneas', 'Descripción'],
    [
        ['Constantes API', '7-9', 'URLs de los endpoints del backend'],
        ['Modal credenciales', '17-36', 'Abrir/cerrar modal y autocompletar'],
        ['Inicio', '39-56', 'DOMContentLoaded, verificación de sesión'],
        ['Portada interactiva', '58-109', 'Botón WhatsApp, tarjeta 3D con mousemove'],
        ['Login / Logout', '112-158', 'iniciarSesion(), cerrarSesion(), fetch a LoginController'],
        ['Productos', '161-211', 'Carga, búsqueda y renderizado de tabla'],
        ['Carrito', '219-287', 'Agregar/quitar items, cambio de cantidad, render y cálculo de IGV'],
        ['Registrar venta', '290-319', 'Envía JSON a VentaController, refresca datos'],
        ['Historial', '321-347', 'Carga y renderiza últimas ventas'],
        ['Limpiar', '349-353', 'Reinicia carrito y formulario'],
    ]
)

doc.add_paragraph(
    'Todas las llamadas AJAX usan fetch() con async/await. No depende de ninguna '
    'biblioteca externa (sin jQuery, sin frameworks).'
)

doc.add_paragraph(
    'Constantes de API:\n'
    '  API_LOGIN    = "../backend/controllers/LoginController.php"\n'
    '  API_PRODUCTO = "../backend/controllers/ProductoController.php"\n'
    '  API_VENTA    = "../backend/controllers/VentaController.php"'
)

doc.add_heading('6.4 Flujo de la Aplicación', level=2)
steps = [
    'El usuario accede a index.html. El DOMContentLoaded ejecuta verificarSesionActiva().',
    'Si hay sesión activa, muestra el panel de ventas. Si no, muestra el login.',
    'El usuario ingresa credenciales y hace clic en "Ingresar". '
    'Se envía POST a LoginController. Si es exitoso, se muestra el panel.',
    'En el panel de ventas, se cargan los productos (GET a ProductoController) '
    'y el historial (GET a VentaController).',
    'El usuario busca productos, los agrega al carrito, define cantidades.',
    'Al hacer clic en "Registrar venta", se envía POST JSON a VentaController.',
    'El backend ejecuta una transacción: crea/reusa cliente, inserta venta, '
    'inserta detalles, descuenta stock. Todo o nada.',
    'Si la venta se registra, se refrescan productos (stock actualizado) e historial.',
    'El usuario puede cerrar sesión con el botón "Cerrar sesión".',
]
for i, s in enumerate(steps, 1):
    doc.add_paragraph(f'{i}. {s}')

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 7. API ENDPOINTS — REFERENCIA RÁPIDA
# ══════════════════════════════════════════════════════════════
doc.add_heading('7. API Endpoints — Referencia Rápida', level=1)
add_simple_table(doc,
    ['Endpoint', 'Método', 'Acción', 'Body / Params', 'Respuesta'],
    [
        ['LoginController.php', 'POST', 'login', 'usuario, password (FormData)', '{ok, mensaje}'],
        ['LoginController.php', 'POST', 'logout', '—', '{ok, mensaje}'],
        ['LoginController.php', 'GET', 'estado', '—', '{ok, usuario?}'],
        ['ProductoController.php', 'GET', 'listar', 'busqueda (query)', '{ok, productos[]}'],
        ['ProductoController.php', 'POST', 'crear', 'nombre, id_categoria, precio, stock', '{ok, mensaje}'],
        ['ProductoController.php', 'POST', 'editar', 'id, nombre, id_categoria, precio, stock', '{ok, mensaje}'],
        ['ProductoController.php', 'POST', 'eliminar', 'id', '{ok, mensaje}'],
        ['VentaController.php', 'POST', 'registrar', 'JSON {cliente, items[]}', '{ok, mensaje}'],
        ['VentaController.php', 'GET', 'listar', '—', '{ok, ventas[]}'],
    ]
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 8. INSTALACIÓN Y DESPLIEGUE
# ══════════════════════════════════════════════════════════════
doc.add_heading('8. Instalación y Despliegue', level=1)

doc.add_heading('8.1 Requisitos', level=2)
add_simple_table(doc,
    ['Software', 'Versión Mínima', 'Descarga'],
    [
        ['XAMPP', '8.x', 'https://www.apachefriends.org/'],
        ['PHP', '8.0', 'Incluido en XAMPP'],
        ['MySQL / MariaDB', '10.x', 'Incluido en XAMPP'],
        ['Navegador web', 'Moderno (Chrome, Edge, Firefox)', '—'],
    ]
)

doc.add_heading('8.2 Pasos de Instalación', level=2)
steps = [
    'Descargar e instalar XAMPP desde https://www.apachefriends.org/.',
    'Clonar o copiar la carpeta coquito_proyecto/ dentro de C:\\xampp\\htdocs\\ (o el directorio htdocs correspondiente).',
    'Iniciar Apache y MySQL desde el Panel de Control de XAMPP.',
    'Abrir phpMyAdmin (http://localhost/phpmyadmin) o la consola MySQL.',
    'Ejecutar el archivo database/bd_coquito.sql para crear la base de datos y cargar los datos iniciales.',
    'Abrir el navegador en http://localhost/coquito_proyecto/frontend/index.html.',
    'Iniciar sesión con las credenciales demo: usuario = "administrador", contraseña = "usuario123".',
]
for i, s in enumerate(steps, 1):
    doc.add_paragraph(f'{i}. {s}')

doc.add_heading('8.3 Configuración Adicional', level=2)
doc.add_paragraph(
    'Si se cambian las credenciales de MySQL, editar backend/config/conexion.php '
    'con los valores correctos. Para producción, se recomienda usar variables de '
    'entorno en lugar de credenciales hardcodeadas.'
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 9. GUÍA DE USO
# ══════════════════════════════════════════════════════════════
doc.add_heading('9. Guía de Uso', level=1)

doc.add_heading('9.1 Inicio de Sesión', level=2)
steps = [
    'Abrir http://localhost/coquito_proyecto/frontend/index.html.',
    'Ingresar usuario y contraseña.',
    'Alternativamente, hacer clic en "Ver credenciales de acceso" y luego "Autocompletar".',
    'Presionar Enter o hacer clic en "Ingresar al sistema".',
]
for i, s in enumerate(steps, 1):
    doc.add_paragraph(f'{i}. {s}')

doc.add_heading('9.2 Registro de una Venta', level=2)
steps = [
    'Escribir el nombre del cliente (opcional).',
    'Buscar productos por nombre en el campo de búsqueda.',
    'Hacer clic en "+ Agregar" para agregar productos al carrito.',
    'Ajustar cantidades con los botones + y − en el carrito.',
    'Verificar el subtotal, IGV (18%) y total calculados automáticamente.',
    'Hacer clic en "Registrar venta".',
    'Confirmar el mensaje de venta exitosa.',
    'El stock se descuenta automáticamente y el historial se actualiza.',
]
for i, s in enumerate(steps, 1):
    doc.add_paragraph(f'{i}. {s}')

doc.add_heading('9.3 Consulta de Historial', level=2)
doc.add_paragraph(
    'El historial de las últimas 20 ventas se carga automáticamente al '
    'ingresar al panel y se actualiza después de cada venta registrada.'
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 10. MANTENIMIENTO Y MEJORAS
# ══════════════════════════════════════════════════════════════
doc.add_heading('10. Mantenimiento y Mejoras', level=1)

doc.add_heading('10.1 Consideraciones de Seguridad', level=2)
bullets = [
    'Las contraseñas se almacenan con bcrypt (password_hash/password_verify).',
    'Las sesiones PHP se usan para autenticación; todas las rutas sensibles '
    'ejecutan requiereSesion().',
    'Se recomienda mover las credenciales de BD a variables de entorno en producción.',
    'El frontend escapa texto con textContent antes de insertarlo en el DOM (XSS prevention).',
]
for b in bullets:
    doc.add_paragraph(b, style='List Bullet')

doc.add_heading('10.2 Posibles Mejoras', level=2)
bullets = [
    'Migrar a PDO con prepared statements (actualmente usa mysqli).',
    'Implementar paginación en el historial de ventas.',
    'Agregar reportes y gráficos de ventas por período.',
    'Implementar roles y permisos granulares (no solo login).',
    'Agregar módulo de gestión de usuarios desde el panel.',
    'Internacionalización (soporte multi-idioma).',
    'Unit tests y integration tests.',
    'Migrar a un framework ligero (Laravel o Slim) para mejor organización.',
    'Agregar facturación electrónica (Sunat, Perú).',
    'Contenedor Docker para facilitar el despliegue.',
]
for b in bullets:
    doc.add_paragraph(b, style='List Bullet')

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# ANEXO A — SCRIPT SQL COMPLETO
# ══════════════════════════════════════════════════════════════
doc.add_heading('Anexo A — Script SQL Completo', level=1)
doc.add_paragraph('A continuación, el script completo de la base de datos bd_coquito:')

add_code_block(doc, '''-- ============================================================
-- Base de datos: bd_coquito
-- Proyecto: Compra tu Coquito - Sistema de Ventas
-- ============================================================

CREATE DATABASE IF NOT EXISTS bd_coquito;
USE bd_coquito;

-- ------------------------------------------------------------
-- Tabla roles
-- ------------------------------------------------------------
CREATE TABLE roles (
    id_rol INT AUTO_INCREMENT PRIMARY KEY,
    nombre_rol VARCHAR(50) NOT NULL
);

-- ------------------------------------------------------------
-- Tabla usuarios
-- ------------------------------------------------------------
CREATE TABLE usuarios (
    id_usuario INT AUTO_INCREMENT PRIMARY KEY,
    nombres VARCHAR(100) NOT NULL,
    apellidos VARCHAR(100) NOT NULL,
    usuario VARCHAR(50) NOT NULL UNIQUE,
    correo VARCHAR(100) NOT NULL UNIQUE,
    password VARCHAR(255) NOT NULL,
    id_rol INT NOT NULL,
    estado TINYINT DEFAULT 1,
    fecha_registro DATETIME DEFAULT CURRENT_TIMESTAMP,
    FOREIGN KEY (id_rol) REFERENCES roles(id_rol)
);

-- ------------------------------------------------------------
-- Tabla categorias
-- ------------------------------------------------------------
CREATE TABLE categorias (
    id_categoria INT AUTO_INCREMENT PRIMARY KEY,
    nombre_categoria VARCHAR(100) NOT NULL,
    estado TINYINT DEFAULT 1
);

-- ------------------------------------------------------------
-- Tabla productos
-- ------------------------------------------------------------
CREATE TABLE productos (
    id_producto INT AUTO_INCREMENT PRIMARY KEY,
    nombre_producto VARCHAR(150) NOT NULL,
    id_categoria INT NOT NULL,
    precio DECIMAL(10,2) NOT NULL,
    stock INT NOT NULL DEFAULT 0,
    estado TINYINT DEFAULT 1,
    fecha_registro DATETIME DEFAULT CURRENT_TIMESTAMP,
    FOREIGN KEY (id_categoria) REFERENCES categorias(id_categoria)
);

-- ------------------------------------------------------------
-- Tabla clientes
-- ------------------------------------------------------------
CREATE TABLE clientes (
    id_cliente INT AUTO_INCREMENT PRIMARY KEY,
    nombre_cliente VARCHAR(150) NOT NULL,
    fecha_registro DATETIME DEFAULT CURRENT_TIMESTAMP
);

-- ------------------------------------------------------------
-- Tabla ventas (cabecera)
-- ------------------------------------------------------------
CREATE TABLE ventas (
    id_venta INT AUTO_INCREMENT PRIMARY KEY,
    id_cliente INT NULL,
    id_usuario INT NOT NULL,
    subtotal DECIMAL(10,2) NOT NULL,
    igv DECIMAL(10,2) NOT NULL,
    total DECIMAL(10,2) NOT NULL,
    fecha_venta DATETIME DEFAULT CURRENT_TIMESTAMP,
    FOREIGN KEY (id_cliente) REFERENCES clientes(id_cliente),
    FOREIGN KEY (id_usuario) REFERENCES usuarios(id_usuario)
);

-- ------------------------------------------------------------
-- Tabla detalle_ventas (items de cada venta)
-- ------------------------------------------------------------
CREATE TABLE detalle_ventas (
    id_detalle INT AUTO_INCREMENT PRIMARY KEY,
    id_venta INT NOT NULL,
    id_producto INT NOT NULL,
    cantidad INT NOT NULL,
    precio_unitario DECIMAL(10,2) NOT NULL,
    subtotal_item DECIMAL(10,2) NOT NULL,
    FOREIGN KEY (id_venta) REFERENCES ventas(id_venta),
    FOREIGN KEY (id_producto) REFERENCES productos(id_producto)
);

-- ------------------------------------------------------------
-- Datos iniciales
-- ------------------------------------------------------------
INSERT INTO roles (nombre_rol) VALUES
('Administrador'),
('Vendedor');

INSERT INTO usuarios (nombres, apellidos, usuario, correo,
    password, id_rol) VALUES
('Administrador', 'Sistema', 'administrador',
 'admin@compratucoquito.com',
 '$2b$10$ivLDo8fsbEOWv1EWtx5h6.43Q18jN1kIAiyySIPIZZdwnrSbNbeFC', 1);

INSERT INTO categorias (nombre_categoria) VALUES
('Papelería'), ('Libros'), ('Útiles'), ('Vinilos');

INSERT INTO productos (nombre_producto, id_categoria, precio, stock) VALUES
('Cuaderno A4 espiral', 1, 8.50, 30),
('Lapicero azul Pilot', 1, 3.00, 50),
('El Principito', 2, 28.00, 10),
('Tijera escolar', 3, 5.50, 0),
('Resaltador x3 colores', 1, 7.00, 20),
('Regla 30cm', 3, 3.00, 15),
('1984 - George Orwell', 2, 35.00, 5),
('Lápices 2B x12', 1, 6.00, 25),
('Vinilo – To Pimp a Butterfly', 4, 89.90, 5),
('Vinilo – Good Kid, M.A.A.D City', 4, 84.90, 5);''')

# ── Guardar ──
output_path = os.path.join(os.path.dirname(__file__), 'Manual_Tecnico_Compra_tu_Coquito.docx')
doc.save(output_path)
print(f'Documento generado: {output_path}')
